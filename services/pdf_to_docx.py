# -*- coding: utf-8 -*-
"""
Servicio de conversion de PDF a DOCX para PDFexport.
Convierte PDF a documento Word intentando preservar formato.
"""

import logging
import io
from pathlib import Path
from typing import Dict, List, Tuple, Optional

import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import config
import models
from utils import file_manager, job_manager

logger = logging.getLogger(__name__)

# Mapeo de calidad de imagenes a DPI
CALIDAD_DPI = {
    'baja': 72,
    'media': 150,
    'alta': 300,
    'original': 0  # 0 significa usar resolucion original
}


def extraer_bloques_texto(pagina: fitz.Page) -> List[Dict]:
    """
    Extrae bloques de texto de una pagina con informacion de formato.

    Args:
        pagina: Pagina de PyMuPDF

    Returns:
        Lista de bloques con texto y formato
    """
    bloques = []

    # Obtener diccionario de texto con informacion detallada
    dict_pagina = pagina.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)

    for bloque in dict_pagina.get("blocks", []):
        if bloque.get("type") == 0:  # Bloque de texto
            for linea in bloque.get("lines", []):
                for span in linea.get("spans", []):
                    texto = span.get("text", "").strip()
                    if texto:
                        bloques.append({
                            'texto': texto,
                            'fuente': span.get("font", ""),
                            'tamano': span.get("size", 12),
                            'color': span.get("color", 0),
                            'flags': span.get("flags", 0),  # negrita, cursiva, etc.
                            'bbox': span.get("bbox", (0, 0, 0, 0))
                        })

    return bloques


def detectar_tabla(pagina: fitz.Page) -> List[Dict]:
    """
    Intenta detectar tablas en una pagina.
    Usa heuristicas basicas para identificar estructuras tabulares.

    Args:
        pagina: Pagina de PyMuPDF

    Returns:
        Lista de tablas detectadas con sus celdas
    """
    tablas = []

    try:
        # PyMuPDF puede encontrar tablas
        tabs = pagina.find_tables()

        for tab in tabs:
            tabla_data = {
                'bbox': tab.bbox,
                'filas': [],
                'num_cols': tab.col_count,
                'num_filas': tab.row_count
            }

            # Extraer contenido de cada celda
            for fila in tab.extract():
                tabla_data['filas'].append(fila)

            tablas.append(tabla_data)

    except Exception as e:
        logger.debug(f"No se pudieron detectar tablas: {e}")

    return tablas


def extraer_imagenes_pagina(pagina: fitz.Page, dpi: int = 150) -> List[Tuple[bytes, str]]:
    """
    Extrae imagenes de una pagina.

    Args:
        pagina: Pagina de PyMuPDF
        dpi: Resolucion para imagenes

    Returns:
        Lista de tuplas (bytes_imagen, extension)
    """
    imagenes = []

    # Obtener lista de imagenes en la pagina
    lista_imagenes = pagina.get_images(full=True)

    for img_info in lista_imagenes:
        try:
            xref = img_info[0]
            imagen_base = pagina.parent.extract_image(xref)

            if imagen_base:
                img_bytes = imagen_base["image"]
                ext = imagen_base["ext"]

                # Si se especifica DPI diferente a original, redimensionar
                if dpi > 0 and dpi != 72:
                    # Usar pixmap para redimensionar
                    pix = fitz.Pixmap(img_bytes)
                    if pix.width > 0 and pix.height > 0:
                        # Calcular nuevo tamano basado en DPI
                        factor = dpi / 72.0
                        nuevo_ancho = int(pix.width * factor)
                        nuevo_alto = int(pix.height * factor)

                        # Limitar tamano maximo
                        if nuevo_ancho > 2000:
                            factor = 2000 / pix.width
                            nuevo_ancho = 2000
                            nuevo_alto = int(pix.height * factor)

                imagenes.append((img_bytes, ext))

        except Exception as e:
            logger.debug(f"Error extrayendo imagen: {e}")

    return imagenes


def es_negrita(flags: int) -> bool:
    """Verifica si el texto es negrita basado en flags."""
    return bool(flags & 2**4)  # Bit 4 indica bold


def es_cursiva(flags: int) -> bool:
    """Verifica si el texto es cursiva basado en flags."""
    return bool(flags & 2**1)  # Bit 1 indica italic


def color_int_a_rgb(color_int: int) -> Tuple[int, int, int]:
    """Convierte color entero a tupla RGB."""
    r = (color_int >> 16) & 0xFF
    g = (color_int >> 8) & 0xFF
    b = color_int & 0xFF
    return (r, g, b)


def agregar_tabla_docx(documento: Document, tabla_data: Dict):
    """
    Agrega una tabla al documento DOCX.

    Args:
        documento: Documento python-docx
        tabla_data: Datos de la tabla detectada
    """
    if not tabla_data['filas']:
        return

    num_filas = len(tabla_data['filas'])
    num_cols = tabla_data['num_cols']

    if num_filas == 0 or num_cols == 0:
        return

    # Crear tabla
    tabla = documento.add_table(rows=num_filas, cols=num_cols)
    tabla.style = 'Table Grid'

    # Llenar celdas
    for i, fila in enumerate(tabla_data['filas']):
        for j, celda_texto in enumerate(fila):
            if j < num_cols:
                celda = tabla.cell(i, j)
                celda.text = str(celda_texto) if celda_texto else ""


def convertir_pdf_a_docx(ruta_pdf: Path, opciones: Dict, trabajo_id: str) -> Path:
    """
    Convierte un PDF a DOCX.

    Args:
        ruta_pdf: Ruta al archivo PDF
        opciones: Opciones de conversion
        trabajo_id: ID del trabajo para progreso

    Returns:
        Ruta al archivo DOCX generado
    """
    preservar_imagenes = opciones.get('preservar_imagenes', True)
    preservar_tablas = opciones.get('preservar_tablas', True)
    preservar_estilos = opciones.get('preservar_estilos', True)
    calidad_imagenes = opciones.get('calidad_imagenes', 'media')

    dpi = CALIDAD_DPI.get(calidad_imagenes, 150)

    job_manager.actualizar_progreso(trabajo_id, 5, "Abriendo documento PDF")

    # Abrir PDF
    doc_pdf = fitz.open(str(ruta_pdf))
    num_paginas = len(doc_pdf)

    # Crear documento Word
    documento = Document()

    # Procesar cada pagina
    for num_pag, pagina in enumerate(doc_pdf):
        progreso = 10 + int((num_pag / num_paginas) * 80)
        job_manager.actualizar_progreso(
            trabajo_id,
            progreso,
            f"Procesando pagina {num_pag + 1} de {num_paginas}"
        )

        # Detectar y agregar tablas primero
        if preservar_tablas:
            tablas = detectar_tabla(pagina)
            for tabla_data in tablas:
                agregar_tabla_docx(documento, tabla_data)
                documento.add_paragraph()  # Espacio despues de tabla

        # Extraer texto con formato
        bloques = extraer_bloques_texto(pagina)

        # Agrupar bloques en parrafos (por posicion Y similar)
        parrafo_actual = documento.add_paragraph()
        y_anterior = -1
        umbral_y = 15  # Pixeles de diferencia para considerar nuevo parrafo

        for bloque in bloques:
            y_actual = bloque['bbox'][1]

            # Si hay cambio significativo en Y, nuevo parrafo
            if y_anterior >= 0 and abs(y_actual - y_anterior) > umbral_y:
                parrafo_actual = documento.add_paragraph()

            # Agregar texto con formato
            run = parrafo_actual.add_run(bloque['texto'] + " ")

            if preservar_estilos:
                # Aplicar tamano de fuente
                run.font.size = Pt(bloque['tamano'])

                # Aplicar negrita/cursiva
                if es_negrita(bloque['flags']):
                    run.bold = True
                if es_cursiva(bloque['flags']):
                    run.italic = True

                # Aplicar color (si no es negro)
                if bloque['color'] != 0:
                    r, g, b = color_int_a_rgb(bloque['color'])
                    run.font.color.rgb = RGBColor(r, g, b)

            y_anterior = y_actual

        # Extraer e insertar imagenes
        if preservar_imagenes:
            imagenes = extraer_imagenes_pagina(pagina, dpi)

            for img_bytes, ext in imagenes:
                try:
                    # Crear stream de imagen
                    img_stream = io.BytesIO(img_bytes)

                    # Agregar imagen al documento
                    # Limitar ancho a 6 pulgadas
                    documento.add_picture(img_stream, width=Inches(6))
                    documento.add_paragraph()  # Espacio despues de imagen

                except Exception as e:
                    logger.warning(f"Error insertando imagen: {e}")

        # Salto de pagina entre paginas del PDF (excepto la ultima)
        if num_pag < num_paginas - 1:
            documento.add_page_break()

    doc_pdf.close()

    job_manager.actualizar_progreso(trabajo_id, 90, "Guardando documento DOCX")

    # Guardar documento
    nombre_base = Path(ruta_pdf).stem
    ruta_docx = config.OUTPUT_FOLDER / f"{trabajo_id}_{nombre_base}.docx"
    documento.save(str(ruta_docx))

    return ruta_docx


def procesar_to_docx(trabajo_id: str, archivo_id: str, parametros: dict) -> dict:
    """
    Procesador principal de conversion PDF a DOCX.
    Esta funcion es llamada por el job_manager.

    Args:
        trabajo_id: ID del trabajo
        archivo_id: ID del archivo a procesar
        parametros: Opciones de conversion

    Returns:
        dict con ruta_resultado y mensaje
    """
    # Obtener archivo
    archivo = models.obtener_archivo(archivo_id)
    if not archivo:
        raise ValueError("Archivo no encontrado")

    ruta_pdf = Path(archivo['ruta_archivo'])
    if not ruta_pdf.exists():
        raise ValueError("Archivo fisico no encontrado")

    job_manager.actualizar_progreso(trabajo_id, 2, "Iniciando conversion a DOCX")

    # Convertir
    ruta_docx = convertir_pdf_a_docx(ruta_pdf, parametros, trabajo_id)

    job_manager.actualizar_progreso(trabajo_id, 95, "Comprimiendo archivo")

    # Crear ZIP
    nombre_base = Path(archivo['nombre_original']).stem
    nombre_docx = f"{nombre_base}.docx"
    nombre_zip = f"{trabajo_id}_{nombre_base}_docx.zip"

    archivos_para_zip = [(str(ruta_docx), nombre_docx)]
    ruta_zip = file_manager.crear_zip(archivos_para_zip, nombre_zip)

    # Limpiar archivo temporal
    if ruta_docx.exists():
        ruta_docx.unlink()

    return {
        'ruta_resultado': str(ruta_zip),
        'mensaje': f'Documento convertido a DOCX exitosamente'
    }


def obtener_preview_docx(archivo_id: str) -> dict:
    """
    Genera informacion de preview para la conversion.
    Retorna info basica del PDF para mostrar al usuario.

    Args:
        archivo_id: ID del archivo

    Returns:
        dict con informacion del documento
    """
    archivo = models.obtener_archivo(archivo_id)
    if not archivo:
        raise ValueError("Archivo no encontrado")

    ruta_pdf = Path(archivo['ruta_archivo'])
    if not ruta_pdf.exists():
        raise ValueError("Archivo fisico no encontrado")

    # Obtener info del PDF
    doc = fitz.open(str(ruta_pdf))

    info = {
        'num_paginas': len(doc),
        'tiene_imagenes': False,
        'tiene_tablas': False,
        'num_imagenes': 0
    }

    # Revisar primeras paginas para detectar contenido
    for i, pagina in enumerate(doc):
        if i >= 5:  # Solo revisar primeras 5 paginas
            break

        # Contar imagenes
        imagenes = pagina.get_images()
        info['num_imagenes'] += len(imagenes)
        if imagenes:
            info['tiene_imagenes'] = True

        # Detectar tablas
        try:
            tablas = pagina.find_tables()
            if tablas:
                info['tiene_tablas'] = True
        except:
            pass

    doc.close()

    return info


# Registrar el procesador en el job_manager
job_manager.registrar_procesador('to-docx', procesar_to_docx)
